"""端到端测试：计算 + 报告生成"""

from pathlib import Path

from fastapi.testclient import TestClient

from ring_knife.main import app

ROOT = Path(__file__).resolve().parent.parent
OUTPUT = ROOT / "tests" / "output_test_report.docx"

SAMPLE_PAYLOAD = {
    "project": {
        "entrust_no": "WT20260001",
        "report_no": "BG20260001",
        "entrust_unit": "测试委托单位",
        "contact": "13800000000",
        "supervision_unit": "测试监理单位",
        "construction_unit": "测试施工单位",
        "project_name": "测试工程",
        "unit_address": "上海市测试路1号",
        "project_address": "工程地址测试",
        "entrust_date": "2026-06-01",
        "project_section": "路基",
        "report_date": "2026-06-09",
        "test_nature": "委托检测",
    },
    "params": {
        "soil_type": "素土",
        "max_dry_density": 1.85,
        "compaction_method": "分层夯实",
        "optimal_moisture": 12.5,
        "standards": ["JTG 3450-2019"],
        "ring_spec": "200cm³",
        "design_requirement": 0.93,
        "sample_name": "回填土",
        "material_type": "素土",
        "test_basis": "JTG 3450-2019",
        "judge_basis": "JTG 3450-2019",
        "result_type": "compaction_coeff",
        "equipment_balance": ["2161730"],
        "equipment_oven": ["8161209"],
        "test_location": "河滨北路370号",
        "remark": "",
    },
    "samples": [
        {
            "sample_no": "1",
            "elevation": "100",
            "sampling_date": "2026-06-08",
            "test_date": "2026-06-09",
            "thickness": "300",
            "ring_sample_mass": 432.5,
            "ring_mass": 200.0,
            "ring_volume": 200.0,
            "boxes": [
                {"box_no": "A1", "box_mass": 15.2, "wet_sample_mass": 45.6, "dry_sample_mass": 42.1},
                {"box_no": "A2", "box_mass": 15.0, "wet_sample_mass": 44.8, "dry_sample_mass": 41.5},
            ],
        },
        {
            "sample_no": "2",
            "elevation": "105",
            "sampling_date": "2026-06-08",
            "test_date": "2026-06-09",
            "thickness": "300",
            "ring_sample_mass": 428.0,
            "ring_mass": 200.0,
            "ring_volume": 200.0,
            "boxes": [
                {"box_no": "B1", "box_mass": 15.1, "wet_sample_mass": 46.0, "dry_sample_mass": 42.3},
                {"box_no": "B2", "box_mass": 15.3, "wet_sample_mass": 45.2, "dry_sample_mass": 41.8},
            ],
        },
    ],
}


def test_calc():
    client = TestClient(app)
    resp = client.post("/api/calc", json={"params": SAMPLE_PAYLOAD["params"], "samples": SAMPLE_PAYLOAD["samples"]})
    assert resp.status_code == 200, resp.text
    data = resp.json()
    assert len(data["results"]) == 2
    r0 = data["results"][0]
    assert r0["wet_mass"] == 232.5
    assert r0["wet_density"] == 1.16
    assert r0["avg_moisture"] is not None
    assert r0["dry_density"] is not None
    assert r0["compaction_coeff"] is not None
    print("calc OK:", r0)


def test_report_generate():
    from docx import Document
    import io

    client = TestClient(app)
    resp = client.post("/api/report/generate", json=SAMPLE_PAYLOAD)
    assert resp.status_code == 200, resp.text
    assert resp.headers["content-type"].startswith("application/vnd.openxmlformats")
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    OUTPUT.write_bytes(resp.content)
    assert OUTPUT.stat().st_size > 10000

    doc = Document(io.BytesIO(resp.content))
    assert SAMPLE_PAYLOAD["project"]["entrust_no"] in doc.paragraphs[2].text
    assert SAMPLE_PAYLOAD["project"]["test_nature"] in doc.paragraphs[3].text
    assert SAMPLE_PAYLOAD["project"]["report_no"] in doc.paragraphs[3].text
    assert "共1页" in doc.paragraphs[3].text
    assert "第一页" in doc.paragraphs[3].text
    table = doc.tables[0]
    design_cell = table.rows[7].cells[1].text
    assert "≥0.93" in design_cell
    assert table.rows[1].cells[0].text.strip() == "监理单位"
    assert "测试监理单位" in table.rows[1].cells[1].text
    assert "测试施工单位" in table.rows[1].cells[6].text
    print(f"report OK: {OUTPUT} ({OUTPUT.stat().st_size} bytes)")


def test_index():
    client = TestClient(app)
    resp = client.get("/")
    assert resp.status_code == 200
    assert "环刀法" in resp.text


if __name__ == "__main__":
    test_calc()
    test_report_generate()
    test_index()
    print("All e2e tests passed.")
