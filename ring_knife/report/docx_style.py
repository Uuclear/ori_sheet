from __future__ import annotations

import re

from docx.document import Document as DocumentObject
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph

# 五号 ≈ 10.5pt
FONT_SIZE = Pt(10.5)
CN_FONT = "仿宋_GB2312"
EN_FONT = "Times New Roman"

_CJK_RE = re.compile(
    r"[\u2e80-\u9fff\u3400-\u4dbf\uf900-\ufaff\u3000-\u303f\uff00-\uffef]"
)


def _is_cjk_char(ch: str) -> bool:
    return bool(_CJK_RE.match(ch))


def _split_text_segments(text: str) -> list[tuple[str, bool]]:
    if not text:
        return []
    segments: list[tuple[str, bool]] = []
    buf: list[str] = []
    is_cjk: bool | None = None
    for ch in text:
        cjk = _is_cjk_char(ch)
        if is_cjk is None:
            is_cjk = cjk
            buf.append(ch)
            continue
        if cjk == is_cjk:
            buf.append(ch)
        else:
            segments.append(("".join(buf), is_cjk))
            buf = [ch]
            is_cjk = cjk
    if buf:
        segments.append(("".join(buf), bool(is_cjk)))
    return segments


def _apply_run_font(run, is_cjk: bool) -> None:
    run.font.size = FONT_SIZE
    run.font.name = EN_FONT
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), EN_FONT)
    r_fonts.set(qn("w:hAnsi"), EN_FONT)
    r_fonts.set(qn("w:cs"), EN_FONT)
    r_fonts.set(qn("w:eastAsia"), CN_FONT if is_cjk else EN_FONT)


def set_paragraph_styled_text(paragraph: Paragraph, text: str) -> None:
    paragraph.text = ""
    if not text:
        return
    for chunk, is_cjk in _split_text_segments(text):
        run = paragraph.add_run(chunk)
        _apply_run_font(run, is_cjk)


def set_cell_text_styled(cell, value: str) -> None:
    text = "" if value is None else str(value)
    cell.text = ""
    set_paragraph_styled_text(cell.paragraphs[0], text)


def style_document(doc: DocumentObject) -> None:
    for paragraph in doc.paragraphs:
        text = paragraph.text
        if text:
            set_paragraph_styled_text(paragraph, text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text
                if text:
                    set_paragraph_styled_text(cell.paragraphs[0], text)
