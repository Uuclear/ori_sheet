from __future__ import annotations

import io
import shutil
import tempfile
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement

from ring_knife.calc.ring_knife import calculate_all
from ring_knife.config import settings
from ring_knife.report.docx_style import set_cell_text_styled, set_paragraph_styled_text, style_document
from ring_knife.schemas.models import CalcRequest, ProjectInfo, RecordParams, ReportRequest

# 环刀报告模板为单页版式
REPORT_TOTAL_PAGES = 1


def _set_cell_text(table, row: int, col: int, value: str) -> None:
    if row < 0 or col < 0:
        return
    if row >= len(table.rows):
        return
    cells = table.rows[row].cells
    if col >= len(cells):
        return
    set_cell_text_styled(cells[col], value)


def _fmt(value: float | None, places: int = 2) -> str:
    if value is None:
        return ""
    return f"{value:.{places}f}"


def _fmt_design_requirement(params: RecordParams) -> str:
    if params.design_requirement is None:
        return ""
    val = _fmt(params.design_requirement)
    if params.result_type == "compaction_percent":
        return f"≥{val}%"
    return f"≥{val}"


def _page_label_cn(page: int) -> str:
    digits = ("零", "一", "二", "三", "四", "五", "六", "七", "八", "九", "十")
    if page <= 0:
        return str(page)
    if page <= 10:
        return digits[page]
    if page < 20:
        return "十" + (digits[page % 10] if page % 10 else "")
    tens, ones = divmod(page, 10)
    text = digits[tens] + "十"
    if ones:
        text += digits[ones]
    return text


def _make_blank_row(ncols: int = 11):
    tr = OxmlElement("w:tr")
    for _ in range(ncols):
        tc = OxmlElement("w:tc")
        tc_pr = OxmlElement("w:tcPr")
        tc.append(tc_pr)
        p = OxmlElement("w:p")
        tc.append(p)
        tr.append(tc)
    return tr


def _set_header_label(table, row: int, left: str, right: str) -> None:
    if row >= len(table.rows):
        return
    cells = table.rows[row].cells
    if cells:
        set_cell_text_styled(cells[0], left)
    if len(cells) > 5:
        set_cell_text_styled(cells[5], right)


def _ensure_project_table_layout(table) -> tuple[int, int]:
    """确保委托信息含监理/施工单位行，返回 (样品数据起始行, 结论行)。"""
    if len(table.rows) < 2:
        return 10, 13
    row1_label = table.rows[1].cells[0].text.strip()
    if "监理单位" in row1_label:
        return 11, 14

    table._tbl.insert(2, _make_blank_row(len(table.rows[0].cells)))
    _set_header_label(table, 1, "监理单位", "施工单位")
    _set_header_label(table, 2, "工程名称", "单位地址")
    return 11, 14


def _fill_project_info(table, project: ProjectInfo, params: RecordParams) -> None:
    design_text = _fmt_design_requirement(params)
    pairs = [
        (0, 1, project.entrust_unit),
        (0, 6, project.contact),
        (1, 1, project.supervision_unit),
        (1, 6, project.construction_unit),
        (2, 1, project.project_name),
        (2, 6, project.unit_address),
        (3, 1, project.project_address),
        (3, 6, project.entrust_date),
        (4, 1, project.project_section),
        (4, 6, project.report_date),
        (5, 1, params.sample_name),
        (5, 6, params.material_type),
        (6, 1, params.ring_spec),
        (6, 6, params.compaction_method),
        (7, 1, design_text),
        (7, 6, _fmt(params.max_dry_density) if params.max_dry_density is not None else ""),
        (8, 1, params.test_location),
        (8, 6, _fmt(params.optimal_moisture) if params.optimal_moisture is not None else ""),
        (9, 1, params.test_basis),
        (9, 6, params.judge_basis),
    ]
    for row, col, text in pairs:
        _set_cell_text(table, row, col, text or "")


def _fill_sample_rows(
    table,
    results,
    params: RecordParams,
    start_row: int,
) -> None:
    for i, result in enumerate(results[:3]):
        row = start_row + i
        _set_cell_text(table, row, 0, result.sample_no)
        _set_cell_text(table, row, 1, result.elevation)
        _set_cell_text(table, row, 2, result.thickness)
        _set_cell_text(table, row, 3, result.sampling_date)
        _set_cell_text(table, row, 4, result.test_date)
        _set_cell_text(table, row, 7, _fmt(result.avg_wet_density or result.wet_density))
        _set_cell_text(table, row, 8, _fmt(result.avg_moisture))
        _set_cell_text(table, row, 9, _fmt(result.avg_dry_density or result.dry_density))
        if params.result_type == "compaction_percent":
            value = result.compaction_percent
        else:
            value = result.compaction_coeff
        _set_cell_text(table, row, 10, _fmt(value))


def _fill_header_paragraphs(doc: Document, project: ProjectInfo) -> None:
    """填充模板页眉：委托编号、检测性质、页码、报告编号。"""
    if len(doc.paragraphs) < 4:
        return
    set_paragraph_styled_text(
        doc.paragraphs[2],
        f"委托编号：{project.entrust_no}",
    )
    page_cn = _page_label_cn(1)
    set_paragraph_styled_text(
        doc.paragraphs[3],
        (
            f"检测性质：{project.test_nature}"
            f"                      共{REPORT_TOTAL_PAGES}页，第{page_cn}页"
            f"                         报告编号：{project.report_no}"
        ),
    )


def _replace_placeholders(doc: Document, mapping: dict[str, str]) -> None:
    for paragraph in doc.paragraphs:
        for key, val in mapping.items():
            token = f"{{{{{key}}}}}"
            if token in paragraph.text:
                set_paragraph_styled_text(paragraph, paragraph.text.replace(token, val))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, val in mapping.items():
                    token = f"{{{{{key}}}}}"
                    if token in cell.text:
                        set_cell_text_styled(cell, cell.text.replace(token, val))


def generate_report(request: ReportRequest) -> bytes:
    template_path = settings.report_template_path
    if not template_path.exists():
        raise FileNotFoundError(f"报告模板不存在: {template_path}")

    calc = calculate_all(CalcRequest(params=request.params, samples=request.samples))

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = Path(tmp.name)
    try:
        shutil.copy(template_path, tmp_path)
        doc = Document(str(tmp_path))

        mapping = _build_placeholder_mapping(request.project, request.params, calc.overall_conclusion)
        _replace_placeholders(doc, mapping)
        _fill_header_paragraphs(doc, request.project)

        if doc.tables:
            table = doc.tables[0]
            sample_start, conclusion_row = _ensure_project_table_layout(table)
            _fill_project_info(table, request.project, request.params)
            _fill_sample_rows(table, calc.results, request.params, sample_start)
            _set_cell_text(table, conclusion_row, 1, calc.overall_conclusion)

        style_document(doc)

        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()
    finally:
        tmp_path.unlink(missing_ok=True)


def _build_placeholder_mapping(
    project: ProjectInfo,
    params: RecordParams,
    conclusion: str,
) -> dict[str, str]:
    design_text = _fmt_design_requirement(params)
    return {
        "委托编号": project.entrust_no,
        "报告编号": project.report_no,
        "委托单位": project.entrust_unit,
        "联系方式": project.contact,
        "监理单位": project.supervision_unit,
        "施工单位": project.construction_unit,
        "工程名称": project.project_name,
        "单位地址": project.unit_address,
        "工程地址": project.project_address,
        "委托日期": project.entrust_date,
        "工程部位": project.project_section,
        "报告日期": project.report_date,
        "检测性质": project.test_nature,
        "样品名称": params.sample_name,
        "材料种类": params.material_type,
        "环刀规格": params.ring_spec,
        "夯实方式": params.compaction_method,
        "设计要求": design_text,
        "最大干密度": _fmt(params.max_dry_density) if params.max_dry_density is not None else "",
        "最优含水率": _fmt(params.optimal_moisture) if params.optimal_moisture is not None else "",
        "检测依据": params.test_basis,
        "判定依据": params.judge_basis,
        "检测结论": conclusion,
    }
